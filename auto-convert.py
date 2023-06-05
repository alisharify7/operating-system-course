import os
import pandas as pd
from docx import Document
import pathlib
from weasyprint import HTML


file_name = 'Lecture-6.docx'
file_path =  pathlib.Path(__file__).parent / file_name

doc = Document(file_path)


html_output_path = './output.html'
with open(html_output_path, 'w', encoding='utf-8') as f:
    f.write('<html><head></head><body>')
    for paragraph in doc.paragraphs:
        f.write(f"<p>{paragraph.text}</p>")
    f.write('</body></html>')

txt_output_path = './output.txt'
with open(txt_output_path, 'w',encoding='utf-8') as f:
    for paragraph in doc.paragraphs:
        f.write(paragraph.text + '\n')


pdf_output_path = './output.pdf'
HTML(filename=txt_output_path).write_pdf(pdf_output_path)

